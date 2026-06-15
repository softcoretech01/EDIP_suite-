import os
import shutil
import logging
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session

from ..database.database import get_db
from ..models import models
from ..auth.auth import get_current_user
from ..vector_db.qdrant_document_service import get_qdrant_doc_service
from ..embeddings.metadata_embedder import MetadataEmbedder

# Import pandas and document libraries dynamically or protect with try/except
try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pypdf
except ImportError:
    pypdf = None

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/uploads",
    tags=["uploads"],
)

embedder = MetadataEmbedder()

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


def df_to_markdown_manual(df: "pd.DataFrame", row_start: int, row_end: int) -> str:
    sub_df = df.iloc[row_start:row_end]
    headers = [str(c) for c in sub_df.columns]
    rows = []
    for _, r in sub_df.iterrows():
        rows.append([str(val) for val in r.values])
    
    # Format table
    header_line = "| " + " | ".join(headers) + " |"
    sep_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    body_lines = []
    for row in rows:
        body_lines.append("| " + " | ".join(row) + " |")
    
    return "\n".join([header_line, sep_line] + body_lines)


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> List[str]:
    """Chunks normal text documents into overlapping segments."""
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end]
        chunks.append(chunk)
        if end == text_len:
            break
        start += (chunk_size - overlap)
    return chunks


@router.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Upload and index client data in Excel, Word, PDF, or text formats.
    """
    filename = file.filename
    ext = os.path.splitext(filename)[1].lower()
    
    allowed_extensions = {".xlsx", ".xls", ".docx", ".pdf", ".txt", ".csv"}
    if ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format. Supported formats: {', '.join(allowed_extensions)}"
        )

    # Save file to uploads directory
    file_id_str = str(models.func.now())  # Keep it simple, or generate unique name
    import time
    timestamp = int(time.time())
    unique_filename = f"{timestamp}_{filename}"
    filepath = os.path.join(UPLOAD_DIR, unique_filename)
    
    try:
        with open(filepath, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        logger.error(f"Failed to save uploaded file: {e}")
        raise HTTPException(status_code=500, detail="Failed to save file on disk.")

    # Parse and extract chunks based on file type
    chunks = []
    chunk_metadatas = []
    file_type = "text"

    try:
        if ext in {".xlsx", ".xls"}:
            file_type = "excel"
            if pd is None:
                raise ImportError("pandas is not installed.")
            
            # Read all sheets
            excel_file = pd.ExcelFile(filepath)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                # Remove completely empty rows/columns
                df = df.dropna(how="all")
                if df.empty:
                    continue
                
                # Chunk by groups of 10 rows
                num_rows = len(df)
                chunk_step = 10
                for start_row in range(0, num_rows, chunk_step):
                    end_row = min(start_row + chunk_step, num_rows)
                    table_md = df_to_markdown_manual(df, start_row, end_row)
                    chunk_text_data = f"File: {filename}\nSheet: {sheet_name}\nRows {start_row + 1} to {end_row} of {num_rows}:\n{table_md}"
                    chunks.append(chunk_text_data)
                    chunk_metadatas.append({
                        "sheet_name": sheet_name,
                        "row_start": start_row,
                        "row_end": end_row,
                        "total_rows": num_rows
                    })

        elif ext == ".csv":
            file_type = "excel"
            if pd is None:
                raise ImportError("pandas is not installed.")
            
            df = pd.read_csv(filepath)
            df = df.dropna(how="all")
            if not df.empty:
                num_rows = len(df)
                chunk_step = 10
                for start_row in range(0, num_rows, chunk_step):
                    end_row = min(start_row + chunk_step, num_rows)
                    table_md = df_to_markdown_manual(df, start_row, end_row)
                    chunk_text_data = f"File: {filename}\nRows {start_row + 1} to {end_row} of {num_rows}:\n{table_md}"
                    chunks.append(chunk_text_data)
                    chunk_metadatas.append({
                        "row_start": start_row,
                        "row_end": end_row,
                        "total_rows": num_rows
                    })

        elif ext == ".docx":
            file_type = "word"
            if docx is None:
                raise ImportError("python-docx is not installed.")
            
            doc = docx.Document(filepath)
            full_text_list = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text_list.append(para.text.strip())
            
            # Read docx tables as well
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text_list.append(row_text)
            
            full_text = "\n\n".join(full_text_list)
            chunks = chunk_text(full_text)
            chunk_metadatas = [{"paragraph_index": i} for i in range(len(chunks))]

        elif ext == ".pdf":
            file_type = "pdf"
            if pypdf is None:
                raise ImportError("pypdf is not installed.")
            
            reader = pypdf.PdfReader(filepath)
            full_text_list = []
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    full_text_list.append(f"--- Page {page_num + 1} ---\n{page_text.strip()}")
            
            full_text = "\n\n".join(full_text_list)
            chunks = chunk_text(full_text)
            chunk_metadatas = [{"page_index_chunk": i} for i in range(len(chunks))]

        else:  # TXT
            file_type = "text"
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                full_text = f.read()
            chunks = chunk_text(full_text)
            chunk_metadatas = [{"chunk_index_txt": i} for i in range(len(chunks))]

    except Exception as e:
        logger.error(f"Error parsing file: {e}")
        # Clean up file from disk
        if os.path.exists(filepath):
            os.remove(filepath)
        raise HTTPException(status_code=400, detail=f"Failed to parse and extract text from the file: {str(e)}")

    if not chunks:
        # Clean up empty file
        if os.path.exists(filepath):
            os.remove(filepath)
        raise HTTPException(status_code=400, detail="The file is empty or no text could be extracted.")

    # Save file record in DB
    try:
        db_file = models.UploadedFile(
            tenant_id=current_user.tenant_id,
            user_id=current_user.id,
            filename=filename,
            file_type=file_type,
            filepath=filepath
        )
        db.add(db_file)
        db.commit()
        db.refresh(db_file)
    except Exception as e:
        logger.error(f"Failed to save file metadata to DB: {e}")
        db.rollback()
        if os.path.exists(filepath):
            os.remove(filepath)
        raise HTTPException(status_code=500, detail="Failed to save file metadata database record.")

    # Embed and index chunks in Qdrant
    try:
        # Compute embeddings in batch
        vectors = embedder.embed_batch(chunks)
        qdrant_doc_service = get_qdrant_doc_service()
        qdrant_doc_service.index_document_chunks(
            tenant_id=current_user.tenant_id,
            user_id=current_user.id,
            file_id=db_file.id,
            filename=filename,
            file_type=file_type,
            chunks=chunks,
            vectors=vectors,
            chunk_metadatas=chunk_metadatas
        )
    except Exception as e:
        logger.error(f"Failed to generate embeddings or index chunks in Qdrant: {e}")
        # Roll back DB record and delete file
        db.delete(db_file)
        db.commit()
        if os.path.exists(filepath):
            os.remove(filepath)
        raise HTTPException(status_code=500, detail=f"Vector database indexing failed: {str(e)}")

    return {
        "id": db_file.id,
        "filename": db_file.filename,
        "file_type": db_file.file_type,
        "created_at": db_file.created_at,
        "chunks_indexed": len(chunks)
    }


@router.get("/")
def list_files(
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    List all uploaded files for the logged-in user.
    """
    files = db.query(models.UploadedFile).filter(
        models.UploadedFile.tenant_id == current_user.tenant_id,
        models.UploadedFile.user_id == current_user.id
    ).order_by(models.UploadedFile.created_at.desc()).all()
    
    return [
        {
            "id": f.id,
            "filename": f.filename,
            "file_type": f.file_type,
            "created_at": f.created_at
        }
        for f in files
    ]


@router.delete("/{file_id}")
def delete_file(
    file_id: int,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(get_current_user)
):
    """
    Delete an uploaded file and clean up its vectors from Qdrant.
    """
    file_record = db.query(models.UploadedFile).filter(
        models.UploadedFile.id == file_id
    ).first()

    if not file_record:
        raise HTTPException(status_code=404, detail="File record not found.")

    if file_record.tenant_id != current_user.tenant_id or file_record.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden: You do not own this file.")

    # Remove points from Qdrant
    try:
        qdrant_doc_service = get_qdrant_doc_service()
        qdrant_doc_service.delete_file_chunks(file_id)
    except Exception as e:
        logger.error(f"Failed to delete Qdrant points for file {file_id}: {e}")

    # Delete local file from disk
    if file_record.filepath and os.path.exists(file_record.filepath):
        try:
            os.remove(file_record.filepath)
        except Exception as e:
            logger.error(f"Failed to remove file from disk: {e}")

    # Remove record from DB
    try:
        db.delete(file_record)
        db.commit()
    except Exception as e:
        logger.error(f"Failed to delete DB record for file {file_id}: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to delete file database record.")

    return {"status": "success", "message": f"File '{file_record.filename}' deleted successfully."}
